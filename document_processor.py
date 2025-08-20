from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import PyPDF2
from docx import Document as DocxDocument
import os

class DocumentProcessor:
    def __init__(self, chunk_size=1500, chunk_overlap=200):
        """
        Initialize document processor with chunking parameters.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            length_function=len
        )
    
    def process_document(self, file_path, file_name):
        """
        Process a document and return chunks.
        
        Args:
            file_path: Path to the document file
            file_name: Original name of the file
            
        Returns:
            List of document chunks with metadata
        """
        # Extract text based on file type
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            text = self._extract_pdf_text(file_path)
        elif file_extension == '.txt':
            text = self._extract_txt_text(file_path)
        elif file_extension == '.docx':
            text = self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Create chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create Document objects with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": file_name,
                    "chunk_id": i + 1,
                    "total_chunks": len(chunks)
                }
            )
            documents.append(doc)
        
        return documents
    
    def _extract_pdf_text(self, file_path):
        """Extract text from PDF file."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
        return text
    
    def _extract_txt_text(self, file_path):
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_docx_text(self, file_path):
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text